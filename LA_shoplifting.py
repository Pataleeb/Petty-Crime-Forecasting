import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import statsmodels.api as sm
from statsmodels.tsa.stattools import adfuller
from statsmodels.graphics.tsaplots import plot_acf, plot_pacf
from sklearn.metrics import mean_absolute_error, mean_squared_error
from docx import Document
from docx.shared import Inches


file_path = "LA_Shoplifting.csv"
shoplifting_data = pd.read_csv(file_path)

# Convert 'DATE OCC' to datetime format
shoplifting_data['DATE OCC'] = pd.to_datetime(shoplifting_data['DATE OCC'], errors='coerce')

# Aggregate shoplifting incidents by month
shoplifting_monthly = shoplifting_data.resample('M', on='DATE OCC').size().to_frame(name='Shoplifting_Count')


def adf_test(series, title):
    result = adfuller(series.dropna())
    print(f"{title}")
    print("ADF Statistic:", result[0])
    print("p-value:", result[1])
    plt.figure(figsize=(10, 4))
    plt.plot(series, label=title)
    plt.legend()
    plt.title(f"{title} - Time Series Plot")
    plt.savefig(f"{title}.png")
    plt.close()

# Check stationarity before differencing
adf_test(shoplifting_monthly['Shoplifting_Count'], "Original Shoplifting Series")

# Apply first differencing
shoplifting_monthly['Differenced'] = shoplifting_monthly['Shoplifting_Count'].diff()

# Check stationarity after differencing
adf_test(shoplifting_monthly['Differenced'], "Differenced Shoplifting Series")

# Plot ACF and PACF
fig, axes = plt.subplots(1, 2, figsize=(16, 6))
plot_acf(shoplifting_monthly['Differenced'].dropna(), lags=20, ax=axes[0])
axes[0].set_title("Autocorrelation Function (ACF) - First Differenced")
plot_pacf(shoplifting_monthly['Differenced'].dropna(), lags=20, ax=axes[1])
axes[1].set_title("Partial Autocorrelation Function (PACF) - First Differenced")
plt.savefig("ACF_PACF.png")
plt.close()


shoplifting_monthly.index = pd.to_datetime(shoplifting_monthly.index)

# Split data into training (up to 2022) and testing (2023)
train = shoplifting_monthly[shoplifting_monthly.index.year <= 2022]
test = shoplifting_monthly[shoplifting_monthly.index.year == 2023]
n_forecast = len(test)

# Fit SARIMA Model
model = sm.tsa.statespace.SARIMAX(
    train['Shoplifting_Count'],
    order=(2,1,1),
    seasonal_order=(1,1,1,12),
    enforce_stationarity=False,
    enforce_invertibility=False
)

results = model.fit()

# Forecast for 2023 using SARIMA
forecast = results.get_forecast(steps=n_forecast)
forecast_series = pd.Series(forecast.predicted_mean.values, index=test.index)

# Compute error metrics for SARIMA
mae = mean_absolute_error(test['Shoplifting_Count'], forecast_series)
rmse = np.sqrt(mean_squared_error(test['Shoplifting_Count'], forecast_series))
mape = np.mean(np.abs((test['Shoplifting_Count'] - forecast_series) / test['Shoplifting_Count'])) * 100


print(f"Mean Absolute Error (MAE): {mae}")
print(f"Root Mean Squared Error (RMSE): {rmse}")
print(f"Mean Absolute Percentage Error (MAPE): {mape}%")


plt.figure(figsize=(12,6))
plt.plot(test['Shoplifting_Count'], label='Actual (2023)', marker='o')
plt.plot(forecast_series, label='Forecast (2023)', linestyle='dashed', color='red', marker='o')
plt.legend()
plt.title("LA Shoplifting Monthly Counts - Forecast vs Actual (2023)")
plt.xlabel("Month")
plt.ylabel("Shoplifting Count")
plt.grid(True)
plt.savefig("SARIMA_Forecast.png")
plt.close()


doc = Document()
doc.add_heading('LA Shoplifting Analysis', level=1)

doc.add_heading('Stationarity Test - Original Data', level=2)
doc.add_picture("Original Shoplifting Series.png", width=Inches(5))

doc.add_heading('Stationarity Test - Differenced Data', level=2)
doc.add_picture("Differenced Shoplifting Series.png", width=Inches(5))

doc.add_heading('ACF and PACF Plots', level=2)
doc.add_picture("ACF_PACF.png", width=Inches(5))

doc.add_heading('SARIMA Forecast vs Actual (2023)', level=2)
doc.add_picture("SARIMA_Forecast.png", width=Inches(5))


doc.add_heading('Error Metrics', level=2)
doc.add_paragraph(f"Mean Absolute Error (MAE): {mae:.2f}")
doc.add_paragraph(f"Root Mean Squared Error (RMSE): {rmse:.2f}")
doc.add_paragraph(f"Mean Absolute Percentage Error (MAPE): {mape:.2f}%")


doc.save("LA_Shoplifting_Analysis.docx")

print("Analysis document saved as 'LA_Shoplifting_Analysis.docx'")

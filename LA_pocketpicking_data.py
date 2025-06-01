import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import statsmodels.api as sm
from statsmodels.tsa.stattools import adfuller
from statsmodels.graphics.tsaplots import plot_acf, plot_pacf
from sklearn.metrics import mean_absolute_error, mean_squared_error
from docx import Document
from docx.shared import Inches


file_path = "LA_Pickpocket.csv"
la_pickpocket_data = pd.read_csv(file_path)

# Convert 'DATE OCC' to datetime format
la_pickpocket_data['DATE OCC'] = pd.to_datetime(la_pickpocket_data['DATE OCC'], errors='coerce')

# Aggregate pocket picking incidents by month
la_pickpocket_monthly = la_pickpocket_data.resample('M', on='DATE OCC').size().to_frame(name='Pickpocket_Count')


def adf_test(series, title):
    result = adfuller(series.dropna())
    print(f"{title}")
    print("ADF Statistic:", result[0])
    print("p-value:", result[1])
    plt.figure(figsize=(10, 4))
    plt.plot(series, label=title)
    plt.legend()
    plt.title(f"{title} - Time Series Plot")
    plt.savefig(f"{title}.png")
    plt.close()

# Check stationarity before differencing
adf_test(la_pickpocket_monthly['Pickpocket_Count'], "Original Pickpocket Series")

# Apply first differencing
la_pickpocket_monthly['Differenced'] = la_pickpocket_monthly['Pickpocket_Count'].diff()

# Check stationarity after differencing
adf_test(la_pickpocket_monthly['Differenced'], "First Differenced Pickpocket Series")

# Plot ACF and PACF
fig, axes = plt.subplots(1, 2, figsize=(16, 6))
plot_acf(la_pickpocket_monthly['Differenced'].dropna(), lags=20, ax=axes[0])
axes[0].set_title("ACF - First Differenced")
plot_pacf(la_pickpocket_monthly['Differenced'].dropna(), lags=20, ax=axes[1])
axes[1].set_title("PACF - First Differenced")
plt.savefig("ACF_PACF.png")
plt.close()


la_pickpocket_monthly.index = pd.to_datetime(la_pickpocket_monthly.index)

# Split data into training (up to 2022) and testing (2023)
train = la_pickpocket_monthly[la_pickpocket_monthly.index.year <= 2022]
test = la_pickpocket_monthly[la_pickpocket_monthly.index.year == 2023]
n_forecast = len(test)

# Fit SARIMA Model (since SARIMAX failed earlier)
model_sarima = sm.tsa.statespace.SARIMAX(
    train['Pickpocket_Count'],
    order=(1,1,1),
    seasonal_order=(1,1,1,12),
    enforce_stationarity=False,
    enforce_invertibility=False
)

results_sarima = model_sarima.fit()

# Forecast for 2023 using SARIMA
forecast_sarima = results_sarima.get_forecast(steps=n_forecast)
forecast_series_sarima = pd.Series(forecast_sarima.predicted_mean.values, index=test.index)

# Compute error metrics for SARIMA
mae_sarima = mean_absolute_error(test['Pickpocket_Count'], forecast_series_sarima)
rmse_sarima = np.sqrt(mean_squared_error(test['Pickpocket_Count'], forecast_series_sarima))
mape_sarima = np.mean(np.abs((test['Pickpocket_Count'] - forecast_series_sarima) / test['Pickpocket_Count'])) * 100

# Plot actual vs forecasted values for 2023 (SARIMA)
plt.figure(figsize=(12,6))
plt.plot(test['Pickpocket_Count'], label='Actual (2023)', marker='o')
plt.plot(forecast_series_sarima, label='Forecast (2023)', linestyle='dashed', color='red', marker='o')
plt.legend()
plt.title("LA Pickpocket Monthly Counts - Forecast vs Actual (2023) (SARIMA)")
plt.xlabel("Month")
plt.ylabel("Pickpocket Count")
plt.grid(True)
plt.savefig("SARIMA_Forecast.png")
plt.close()


doc = Document()
doc.add_heading('LA Pickpocket Analysis', level=1)

doc.add_heading('Stationarity Test - Original Data', level=2)
doc.add_picture("Original Pickpocket Series.png", width=Inches(5))

doc.add_heading('Stationarity Test - Differenced Data', level=2)
doc.add_picture("First Differenced Pickpocket Series.png", width=Inches(5))

doc.add_heading('ACF and PACF Plots', level=2)
doc.add_picture("ACF_PACF.png", width=Inches(5))

doc.add_heading('SARIMA Forecast vs Actual (2023)', level=2)
doc.add_picture("SARIMA_Forecast.png", width=Inches(5))


doc.add_heading('Error Metrics', level=2)
doc.add_paragraph(f"Mean Absolute Error (MAE): {mae_sarima:.2f}")
doc.add_paragraph(f"Root Mean Squared Error (RMSE): {rmse_sarima:.2f}")
doc.add_paragraph(f"Mean Absolute Percentage Error (MAPE): {mape_sarima:.2f}%")


doc.save("LA_Pickpocket_Analysis.docx")

print("Analysis document saved as 'LA_Pickpocket_Analysis.docx'")

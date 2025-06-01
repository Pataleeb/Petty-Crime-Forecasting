import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import statsmodels.api as sm
from statsmodels.tsa.stattools import adfuller
from statsmodels.graphics.tsaplots import plot_acf, plot_pacf
from sklearn.metrics import mean_absolute_error, mean_squared_error
from docx import Document
from docx.shared import Inches


file_path = "LA_Burglary.csv"
la_burglary_data = pd.read_csv(file_path)

# Convert 'DATE OCC' to datetime format
la_burglary_data['DATE OCC'] = pd.to_datetime(la_burglary_data['DATE OCC'], errors='coerce')

# Convert 'Crm Cd Desc' to binary (Burglary = 0, Burglary from Vehicle = 1)
la_burglary_data['Burglary_Type'] = (la_burglary_data['Crm Cd Desc'] == 'BURGLARY FROM VEHICLE').astype(int)

# Aggregate by month
la_burglary_exog = la_burglary_data.resample('M', on='DATE OCC').agg({'Burglary_Type': 'sum'})
la_burglary_exog['Burglary_Count'] = la_burglary_data.resample('M', on='DATE OCC').size()

# Convert index to datetime
la_burglary_exog.index = pd.to_datetime(la_burglary_exog.index)

# Ensure we exclude 2024 data and only focus on predicting 2023
la_burglary_exog = la_burglary_exog[la_burglary_exog.index.year <= 2023]

# Define exogenous variables used
exogenous_vars = ['Burglary_Type']

# Split data into training (up to 2022) and testing (2023)
train = la_burglary_exog[la_burglary_exog.index.year <= 2022]
test = la_burglary_exog[la_burglary_exog.index.year == 2023]

# Separate exogenous variables for training and testing
train_exog = train[exogenous_vars]
test_exog = test[exogenous_vars]


def adf_test(series, title):
    result = adfuller(series.dropna())
    print(f"{title}")
    print("ADF Statistic:", result[0])
    print("p-value:", result[1])
    plt.figure(figsize=(10, 4))
    plt.plot(series, label=title)
    plt.legend()
    plt.title(f"{title} - Time Series Plot")
    plt.savefig(f"{title}.png")
    plt.close()

# Check stationarity before differencing
adf_test(la_burglary_exog['Burglary_Count'], "Original Burglary Series")

# Apply first differencing
la_burglary_exog['Differenced'] = la_burglary_exog['Burglary_Count'].diff()

# Check stationarity after differencing
adf_test(la_burglary_exog['Differenced'], "Differenced Burglary Series")

# Plot ACF and PACF
fig, axes = plt.subplots(1, 2, figsize=(16, 6))
plot_acf(la_burglary_exog['Differenced'].dropna(), lags=20, ax=axes[0])
axes[0].set_title("Autocorrelation Function (ACF) - First Differenced")
plot_pacf(la_burglary_exog['Differenced'].dropna(), lags=20, ax=axes[1])
axes[1].set_title("Partial Autocorrelation Function (PACF) - First Differenced")
plt.savefig("ACF_PACF.png")
plt.close()

# Fit SARIMAX Model (1,1,1)(1,1,1,12) with exogenous variable
model = sm.tsa.statespace.SARIMAX(
    train['Burglary_Count'],
    exog=train_exog,
    order=(1,1,1),
    seasonal_order=(1,1,1,12),
    enforce_stationarity=False,
    enforce_invertibility=False
)

results = model.fit()

# Forecast for 2023 using exogenous variable
n_forecast = len(test)
forecast = results.get_forecast(steps=n_forecast, exog=test_exog)
forecast_series = pd.Series(forecast.predicted_mean.values, index=test.index)

# Compute error metrics
mae = mean_absolute_error(test['Burglary_Count'], forecast_series)
rmse = np.sqrt(mean_squared_error(test['Burglary_Count'], forecast_series))
mape = np.mean(np.abs((test['Burglary_Count'] - forecast_series) / test['Burglary_Count'])) * 100


print(f"Mean Absolute Error (MAE): {mae}")
print(f"Root Mean Squared Error (RMSE): {rmse}")
print(f"Mean Absolute Percentage Error (MAPE): {mape}%")


plt.figure(figsize=(12,6))
plt.plot(test['Burglary_Count'], label='Actual (2023)', marker='o')
plt.plot(forecast_series, label='Forecast (2023)', linestyle='dashed', color='red', marker='o')
plt.legend()
plt.title("LA Burglary Monthly Counts - Forecast vs Actual (2023) (SARIMAX)")
plt.xlabel("Month")
plt.ylabel("Burglary Count")
plt.grid(True)
plt.savefig("SARIMAX_Forecast.png")
plt.close()


doc = Document()
doc.add_heading('LA Burglary Analysis', level=1)

doc.add_heading('Stationarity Test - Original Data', level=2)
doc.add_picture("Original Burglary Series.png", width=Inches(5))

doc.add_heading('Stationarity Test - Differenced Data', level=2)
doc.add_picture("Differenced Burglary Series.png", width=Inches(5))

doc.add_heading('ACF and PACF Plots', level=2)
doc.add_picture("ACF_PACF.png", width=Inches(5))

doc.add_heading('SARIMAX Forecast vs Actual (2023)', level=2)
doc.add_picture("SARIMAX_Forecast.png", width=Inches(5))

# Add error metrics
doc.add_heading('Error Metrics', level=2)
doc.add_paragraph(f"Mean Absolute Error (MAE): {mae:.2f}")
doc.add_paragraph(f"Root Mean Squared Error (RMSE): {rmse:.2f}")
doc.add_paragraph(f"Mean Absolute Percentage Error (MAPE): {mape:.2f}%")

# Save document
doc.save("LA_Burglary_Analysis.docx")

print("Analysis document saved as 'LA_Burglary_Analysis.docx'")

###Anomaly detection
# Extract actual and predicted values for Jan, July, Oct, Nov, and Dec 2023
selected_months = ["2023-01", "2023-07", "2023-10", "2023-11", "2023-12"]
selected_data = la_burglary_exog.loc[la_burglary_exog.index.strftime('%Y-%m').isin(selected_months), ['Burglary_Count']]
selected_data['Predicted_Burglary_Count'] = forecast_series.loc[selected_data.index]

# Compute Z-score for anomaly detection
selected_data['Z_Score'] = (selected_data['Burglary_Count'] - selected_data['Predicted_Burglary_Count']) / selected_data['Burglary_Count'].std()
selected_data['Anomaly'] = selected_data['Z_Score'].abs() > 2

# Create Word document
doc = Document()
doc.add_heading('LA Burglary Analysis', level=1)

doc.add_heading('Anomaly Detection Results (Z-Scores)', level=2)

table = doc.add_table(rows=1, cols=4)

table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Month'
hdr_cells[1].text = 'Actual Burglary Count'
hdr_cells[2].text = 'Predicted Burglary Count'
hdr_cells[3].text = 'Z-Score'

# Add rows to table
for index, row in selected_data.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(index.strftime('%Y-%m'))
    row_cells[1].text = str(row['Burglary_Count'])
    row_cells[2].text = str(round(row['Predicted_Burglary_Count'], 2))
    row_cells[3].text = str(round(row['Z_Score'], 4))

doc.add_heading('Error Metrics', level=2)
mae = np.mean(np.abs(test['Burglary_Count'] - forecast_series))
rmse = np.sqrt(np.mean((test['Burglary_Count'] - forecast_series) ** 2))
mape = np.mean(np.abs((test['Burglary_Count'] - forecast_series) / test['Burglary_Count'])) * 100

doc.add_paragraph(f"Mean Absolute Error (MAE): {mae:.2f}")
doc.add_paragraph(f"Root Mean Squared Error (RMSE): {rmse:.2f}")
doc.add_paragraph(f"Mean Absolute Percentage Error (MAPE): {mape:.2f}%")

# Save document
doc.save("LA_Burglary_Anomaly_Analysis.docx")

print("Analysis document saved as 'LA_Burglary_Analysis.docx'")
